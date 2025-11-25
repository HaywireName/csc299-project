import os
import uuid
import shutil
import time
from datetime import datetime
from tabulate import tabulate
from pypdf import PdfReader
from docx import Document
from tqdm import tqdm
from openai import OpenAI
from core.errors import PDFNotFoundError, InvalidInputError, APIError, StorageError
from core.utils import confirm_action, format_success, format_error, format_warning, print_progress_bar


class DocumentManager:
    """Manages document operations including PDF, DOCX, and TXT file handling.
    
    The DocumentManager handles document ingestion, text extraction, caching,
    searching, and AI-powered summarization. Supports PDF, DOCX, and TXT formats
    with automatic metadata extraction and full-text search capabilities.
    
    Attributes:
        data_manager: Data persistence manager for document metadata.
        registry: Command registry for registering document commands.
        data_dir: Base directory for document storage.
        pdfs_dir: Directory for PDF files.
        docx_dir: Directory for DOCX files.
        txt_dir: Directory for TXT files.
        cache_dir: Directory for cached extracted text.
        documents: List of document metadata dictionaries.
        session_cost: Cumulative OpenAI API cost for the session.
        openai_client: OpenAI API client instance.
    """
    
    def __init__(self, data_manager, registry):
        """Initialize DocumentManager with dependencies.
        
        Sets up directory structure for document storage, loads existing
        document metadata, initializes OpenAI client, and registers commands.
        
        Args:
            data_manager: Data persistence manager instance.
            registry: Command registry instance for registering commands.
        """
        self.data_manager = data_manager
        self.registry = registry
        self.data_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'docs')
        self.pdfs_dir = os.path.join(self.data_dir, 'pdfs')
        self.docx_dir = os.path.join(self.data_dir, 'docx')
        self.txt_dir = os.path.join(self.data_dir, 'txt')
        self.cache_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'doc_cache')
        
        # Ensure directories exist
        os.makedirs(self.pdfs_dir, exist_ok=True)
        os.makedirs(self.docx_dir, exist_ok=True)
        os.makedirs(self.txt_dir, exist_ok=True)
        os.makedirs(self.cache_dir, exist_ok=True)
        
        self.documents = self._load_documents()
        self.session_cost = 0.0  # Track cumulative cost for the session
        self.openai_client = None
        self._init_openai_client()
        self._register_commands()

    def _init_openai_client(self):
        """Initialize OpenAI client with API key from environment.
        
        Attempts to create an OpenAI client using the API key from the
        OPENAI_API_KEY environment variable. If the key is not found,
        AI summarization features will be disabled.
        """
        try:
            api_key = os.environ.get('OPENAI_API_KEY')
            if api_key:
                self.openai_client = OpenAI(api_key=api_key)
            else:
                print("Warning: OPENAI_API_KEY not found. AI summarization will be disabled.")
        except Exception as e:
            print(f"Warning: Failed to initialize OpenAI client: {e}")

    def _load_documents(self):
        """Load document metadata from storage.
        
        Returns:
            list: List of document metadata dictionaries, or empty list if none exist.
        """
        docs = self.data_manager.load("docs_metadata.json")
        return docs if docs else []

    def _save_documents(self):
        """Save document metadata to storage.
        
        Persists the complete list of document metadata to docs_metadata.json.
        """
        self.data_manager.save("docs_metadata.json", self.documents)

    def _generate_unique_filename(self, original_name):
        """Generate a unique filename with UUID prefix.
        
        Creates a filename in the format: {uuid}_originalname.ext to prevent
        naming conflicts when storing documents.
        
        Args:
            original_name: Original filename with extension.
        
        Returns:
            str: Unique filename with 8-character UUID prefix.
        """
        name, ext = os.path.splitext(original_name)
        unique_id = str(uuid.uuid4())[:8]
        return f"{unique_id}_{original_name}"

    def _extract_title_from_text(self, text, max_length=100):
        """Extract a meaningful title from text content.
        
        Finds the first non-empty, substantial line from text that could serve
        as a title. Filters out very short lines and common non-title patterns.
        
        Args:
            text: Text content to extract title from.
            max_length: Maximum length for the title. Defaults to 100.
        
        Returns:
            str or None: Extracted title, or None if no suitable title found.
        """
        if not text:
            return None
        
        lines = text.split('\n')
        for line in lines:
            # Clean up the line
            cleaned = line.strip()
            
            # Skip empty lines, very short lines, or lines that look like page numbers/dates
            if not cleaned:
                continue
            if len(cleaned) < 3:
                continue
            if cleaned.isdigit():  # Skip pure numbers (page numbers)
                continue
            
            # Check if line looks like a date pattern and skip
            if any(pattern in cleaned.lower() for pattern in ['page ', 'chapter ', '---', '===']):
                continue
            
            # This looks like a good title candidate
            title = cleaned[:max_length]
            if len(cleaned) > max_length:
                title += "..."
            return title
        
        return None

    def _extract_pdf_metadata(self, pdf_path):
        """Extract metadata from PDF file using pypdf.
        
        Extracts page count, title (from metadata or content), and preview
        text from the first page of the PDF. If metadata title is generic or
        missing, extracts title from the first page content.
        
        Args:
            pdf_path: Path to the PDF file.
        
        Returns:
            dict: Dictionary containing:
                - page_count (int): Number of pages in the PDF.
                - title (str or None): PDF title from metadata or content.
                - preview (str): First 500 characters from first page.
        
        Raises:
            ValueError: If PDF cannot be read or metadata extraction fails.
        """
        try:
            reader = PdfReader(pdf_path)
            page_count = len(reader.pages)
            
            # Extract text from first page for preview and title extraction
            preview_text = ""
            if page_count > 0:
                preview_text = reader.pages[0].extract_text()
            
            # Try to get title from metadata first
            title = None
            if reader.metadata and reader.metadata.title:
                metadata_title = reader.metadata.title.strip()
                # Check if metadata title is meaningful (not generic)
                if metadata_title and metadata_title.lower() not in ['(anonymous)', 'untitled', 'document', 'pdf']:
                    title = metadata_title
            
            # If no good title from metadata, extract from content
            if not title:
                title = self._extract_title_from_text(preview_text)
            
            return {
                "page_count": page_count,
                "title": title,
                "preview": preview_text[:500] if preview_text else ""
            }
        except Exception as e:
            raise ValueError(f"Failed to extract PDF metadata: {e}")

    def _extract_docx_metadata(self, docx_path):
        """Extract metadata from DOCX file.
        
        Extracts paragraph count (as proxy for page count), title from
        core properties or content, and preview text from first few paragraphs.
        If core properties title is generic, extracts title from document content.
        
        Args:
            docx_path: Path to the DOCX file.
        
        Returns:
            dict: Dictionary containing:
                - page_count (int): Number of paragraphs in the document.
                - title (str or None): Document title from properties or content.
                - preview (str): First 500 characters from first paragraphs.
        
        Raises:
            ValueError: If DOCX cannot be read or metadata extraction fails.
        """
        try:
            doc = Document(docx_path)
            
            # Count paragraphs as "pages"
            paragraph_count = len(doc.paragraphs)
            
            # Extract text from first few paragraphs for preview
            preview_text = ""
            for para in doc.paragraphs[:10]:  # First 10 paragraphs
                if para.text.strip():  # Only include non-empty paragraphs
                    preview_text += para.text + "\n"
            
            # Try to get title from core properties
            title = None
            if doc.core_properties.title:
                properties_title = doc.core_properties.title.strip()
                # Check if title is meaningful (not generic)
                if properties_title and properties_title.lower() not in [
                    'word document', 'document', 'untitled', 'doc', 'docx'
                ]:
                    title = properties_title
            
            # If no good title from properties, extract from content
            if not title:
                title = self._extract_title_from_text(preview_text)
            
            return {
                "page_count": paragraph_count,
                "title": title,
                "preview": preview_text[:500] if preview_text else ""
            }
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX metadata: {e}")

    def _extract_txt_metadata(self, txt_path):
        """Extract metadata from TXT file.
        
        Extracts line count, uses first meaningful line as title, and creates
        a preview from the beginning of the file.
        
        Args:
            txt_path: Path to the TXT file.
        
        Returns:
            dict: Dictionary containing:
                - page_count (int): Number of lines in the file.
                - title (str or None): First meaningful line as title.
                - preview (str): First 500 characters of the file.
        
        Raises:
            ValueError: If TXT file cannot be read or is invalid.
        """
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            line_count = len(lines)
            
            # Extract title from content
            title = self._extract_title_from_text(content)
            
            return {
                "page_count": line_count,
                "title": title,
                "preview": content[:500] if content else ""
            }
        except Exception as e:
            raise ValueError(f"Failed to extract TXT metadata: {e}")

    def _generate_doc_id(self):
        """Generate a unique numeric ID for a document.
        
        Returns:
            str: Sequential numeric ID as a string (1, 2, 3, ...).
        """
        return str(len(self.documents) + 1)
    
    def _reindex_documents(self):
        """Reindex all documents with sequential IDs."""
        for idx, doc in enumerate(self.documents, start=1):
            doc['id'] = str(idx)

    def add_doc(self, source_path):
        """Add a document (PDF, DOCX, or TXT) to the system.
        
        Validates file type, copies the file to the appropriate storage directory,
        extracts metadata, and creates a document record. Supported formats:
        .pdf, .docx, .txt.
        
        Args:
            source_path: Path to the source document file.
        
        Returns:
            dict: Document metadata dictionary containing id, filename, filepath,
                extension, title, page_count, dates, and preview.
        
        Raises:
            FileNotFoundError: If the source file doesn't exist.
            ValueError: If file type is not supported.
            PermissionError: If file cannot be copied to storage.
        """
        # Validate file exists
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"File not found: {source_path}")
        
        # Get file extension
        filename = os.path.basename(source_path)
        name, ext = os.path.splitext(filename)
        ext_lower = ext.lower()
        
        # Validate extension
        if ext_lower not in ['.pdf', '.docx', '.txt']:
            raise ValueError(f"Unsupported file type: {ext}. Only PDF, DOCX, and TXT files are supported.")
        
        # Determine target directory
        if ext_lower == '.pdf':
            target_dir = self.pdfs_dir
        elif ext_lower == '.docx':
            target_dir = self.docx_dir
        else:  # .txt
            target_dir = self.txt_dir
        
        # Generate unique filename
        unique_filename = self._generate_unique_filename(filename)
        target_path = os.path.join(target_dir, unique_filename)
        
        # Copy file
        try:
            shutil.copy2(source_path, target_path)
        except PermissionError:
            raise PermissionError(f"Permission denied: Cannot copy file to {target_path}")
        except Exception as e:
            raise Exception(f"Failed to copy file: {e}")
        
        # Extract metadata based on file type
        try:
            if ext_lower == '.pdf':
                metadata = self._extract_pdf_metadata(target_path)
            elif ext_lower == '.docx':
                metadata = self._extract_docx_metadata(target_path)
            else:  # .txt
                metadata = self._extract_txt_metadata(target_path)
        except Exception as e:
            # Clean up copied file if metadata extraction fails
            os.remove(target_path)
            raise e
        
        # Create document record
        now = datetime.now().strftime("%m-%d-%YT%H:%M:%S.%f")
        doc = {
            "id": self._generate_doc_id(),
            "filename": filename,
            "filepath": target_path,
            "extension": ext_lower,
            "title": metadata['title'] or name,  # Use filename if no title
            "page_count": metadata['page_count'],
            "added_date": now,
            "last_accessed": now,
            "summary": None,
            "preview": metadata['preview']
        }
        
        self.documents.append(doc)
        self._reindex_documents()
        self._save_documents()
        
        return doc

    def list_docs(self):
        """List all documents sorted by ID (1 to n).
        
        Returns:
            list: List of document metadata dictionaries sorted by ID.
        """
        return sorted(self.documents, key=lambda d: int(d['id']))

    def remove_doc(self, doc_id):
        """Remove a document by ID with confirmation.
        
        Prompts for user confirmation before permanently deleting both
        the document file and its metadata.
        
        Args:
            doc_id: Full or partial document ID.
        
        Raises:
            PDFNotFoundError: If document with the given ID is not found.
            StorageError: If the document file cannot be deleted from storage.
        """
        doc = self.get_doc(doc_id)  # This will raise PDFNotFoundError if not found
        
        # Confirm deletion
        doc_name = doc['filename']
        if not confirm_action(f"Delete document '{doc_name}'? (yes/no):", require_yes=False):
            print(format_warning("Deletion cancelled"))
            return
        
        # Delete file
        try:
            if os.path.exists(doc['filepath']):
                os.remove(doc['filepath'])
        except Exception as e:
            raise StorageError(
                f"Failed to delete file: {str(e)}",
                filepath=doc['filepath'],
                operation="delete"
            )
        
        # Remove from list
        self.documents.remove(doc)
        self._reindex_documents()
        self._save_documents()
        print(format_success(f"Document deleted: {doc_name}"))

    def get_doc(self, doc_id):
        """Retrieve a document by full or partial ID.
        
        Searches for a document where the ID starts with the provided string.
        
        Args:
            doc_id: Full or partial document ID to search for.
        
        Returns:
            dict: Document metadata dictionary.
        
        Raises:
            PDFNotFoundError: If no document with matching ID is found.
                Includes list of available document IDs in the error.
        """
        for doc in self.documents:
            if doc['id'].startswith(str(doc_id)):
                return doc
        
        # Document not found - raise error with available IDs
        available_ids = [doc['id'][:8] for doc in self.documents]
        raise PDFNotFoundError(doc_id, available_ids)

    def update_last_accessed(self, doc_id):
        """Update the last_accessed timestamp for a document.
        
        Updates the last_accessed field to the current timestamp and saves
        the document metadata.
        
        Args:
            doc_id: Full or partial document ID.
        """
        doc = self.get_doc(doc_id)
        if doc:
            doc['last_accessed'] = datetime.now().strftime("%m-%d-%YT%H:%M:%S.%f")
            self._save_documents()

    def _get_cache_path(self, doc_id, page_num=None):
        """Get the cache file path for a document.
        
        Constructs the path for cached extracted text. Supports both full
        document caching and individual page caching for PDFs.
        
        Args:
            doc_id: Document ID.
            page_num: Page number for individual page cache (None for full
                document cache). Defaults to None.
        
        Returns:
            str: Absolute path to the cache file.
        """
        if page_num is None:
            filename = f"{doc_id}_full.txt"
        else:
            filename = f"{doc_id}_page{page_num}.txt"
        return os.path.join(self.cache_dir, filename)

    def _extract_full_text(self, doc_path, extension, page_count):
        """Extract all text from a document.
        
        Extracts complete text content from PDF, DOCX, or TXT files. Shows
        progress bar for large PDFs (>10 pages).
        
        Args:
            doc_path: Path to the document file.
            extension: File extension (.pdf, .docx, or .txt).
            page_count: Number of pages/paragraphs/lines for context.
        
        Returns:
            str: Complete extracted text from the document.
        
        Raises:
            ValueError: If text extraction fails.
        """
        try:
            if extension == '.pdf':
                reader = PdfReader(doc_path)
                text = ""
                total_pages = len(reader.pages)
                
                # Show progress for large PDFs
                if total_pages > 10:
                    print(f"Extracting {total_pages} pages...")
                    for i in tqdm(range(total_pages), desc="Progress"):
                        page_text = reader.pages[i].extract_text()
                        text += f"\n--- Page {i + 1} ---\n{page_text}\n"
                else:
                    for i in range(total_pages):
                        page_text = reader.pages[i].extract_text()
                        text += f"\n--- Page {i + 1} ---\n{page_text}\n"
                
                return text
                
            elif extension == '.docx':
                doc = Document(doc_path)
                text = ""
                for i, para in enumerate(doc.paragraphs, 1):
                    text += f"{para.text}\n"
                return text
                
            else:  # .txt
                with open(doc_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            raise ValueError(f"Failed to extract text: {e}")

    def extract_text(self, doc_id, page_num=None):
        """Extract text from a document with caching.
        
        Extracts text from a complete document or a specific page (PDFs only).
        Uses cached version if available to improve performance. Updates the
        last accessed timestamp.
        
        Args:
            doc_id: Document ID.
            page_num: Page number to extract (PDF only). None extracts full
                document. Defaults to None.
        
        Returns:
            tuple: (text (str), word_count (int)) containing the extracted text
                and word count.
        
        Raises:
            ValueError: If document not found, page extraction requested for
                non-PDF, or invalid page number.
        """
        doc = self.get_doc(doc_id)
        if not doc:
            raise ValueError(f"Document with ID {doc_id} not found.")
        
        # Check cache first
        cache_path = self._get_cache_path(doc_id, page_num)
        if os.path.exists(cache_path):
            with open(cache_path, 'r', encoding='utf-8') as f:
                text = f.read()
            word_count = len(text.split())
            self.update_last_accessed(doc_id)
            return text, word_count
        
        # Extract text
        if page_num is None:
            # Extract full document
            text = self._extract_full_text(doc['filepath'], doc['extension'], doc['page_count'])
        else:
            # Extract single page (PDF only)
            if doc['extension'] != '.pdf':
                raise ValueError("Page extraction only supported for PDF files.")
            
            reader = PdfReader(doc['filepath'])
            if page_num < 1 or page_num > len(reader.pages):
                raise ValueError(f"Invalid page number. Document has {len(reader.pages)} pages.")
            
            text = reader.pages[page_num - 1].extract_text()
        
        # Cache the text
        try:
            with open(cache_path, 'w', encoding='utf-8') as f:
                f.write(text)
        except Exception as e:
            print(f"Warning: Could not cache text: {e}")
        
        # Calculate word count
        word_count = len(text.split())
        
        # Update metadata
        if page_num is None:
            doc['word_count'] = word_count
            self._save_documents()
        
        self.update_last_accessed(doc_id)
        
        return text, word_count

    def search_docs(self, query):
        """Search all documents for a query string.
        
        Performs case-insensitive full-text search across all documents.
        Uses cached text when available. Extracts context around matches
        and identifies page numbers for PDF documents.
        
        Args:
            query: Search query string.
        
        Returns:
            list: List of tuples (doc, page_num, context) where:
                - doc (dict): Document metadata dictionary.
                - page_num (int or None): Page number for PDFs, None for others.
                - context (str): Text excerpt with query highlighted using **.
        """
        if not query:
            return []
        
        query_lower = query.lower()
        results = []
        
        print(f"Searching across {len(self.documents)} document(s)...")
        
        for doc in self.documents:
            try:
                # Extract full text (uses cache if available)
                text, _ = self.extract_text(doc['id'])
                
                # Search for query (case-insensitive)
                text_lower = text.lower()
                index = text_lower.find(query_lower)
                
                while index != -1:
                    # Extract context (100 chars before and after)
                    start = max(0, index - 100)
                    end = min(len(text), index + len(query) + 100)
                    context = text[start:end]
                    
                    # Add ellipsis if truncated
                    if start > 0:
                        context = "..." + context
                    if end < len(text):
                        context = context + "..."
                    
                    # Highlight query in context
                    context_lower = context.lower()
                    query_start = context_lower.find(query_lower)
                    if query_start != -1:
                        highlighted = (
                            context[:query_start] +
                            "**" + context[query_start:query_start + len(query)] + "**" +
                            context[query_start + len(query):]
                        )
                    else:
                        highlighted = context
                    
                    # Determine page number for PDFs
                    page_num = None
                    if doc['extension'] == '.pdf':
                        # Count page markers before this match
                        page_markers = text[:index].count('--- Page ')
                        page_num = page_markers if page_markers > 0 else 1
                    
                    results.append((doc, page_num, highlighted))
                    
                    # Find next occurrence
                    index = text_lower.find(query_lower, index + 1)
                    
            except Exception as e:
                print(f"Warning: Could not search document {doc['id']}: {e}")
                continue
        
        return results

    def _chunk_text(self, text, chunk_size=3000):
        """Split text into chunks of approximately chunk_size words.
        
        Divides long text into manageable chunks for AI processing.
        
        Args:
            text: Text to split.
            chunk_size: Approximate number of words per chunk. Defaults to 3000.
        
        Returns:
            list: List of text chunk strings.
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
        
        return chunks

    def _summarize_text(self, text, max_words, purpose="main"):
        """Summarize text using OpenAI API.
        
        Generates AI-powered summary using gpt-4o-mini. Adjusts prompt based
        on the purpose (main document, chunk, or final synthesis). Implements
        retry logic and tracks costs.
        
        Args:
            text: Text to summarize.
            max_words: Maximum words for the summary.
            purpose: Purpose of summary - "main" for complete documents,
                "chunk" for document sections, "final" for synthesizing
                chunk summaries. Defaults to "main".
        
        Returns:
            tuple: (summary (str), cost (float)) containing the generated
                summary and API cost.
        
        Raises:
            ValueError: If OpenAI client not initialized, API call fails,
                or rate limits exceeded.
        """
        if not self.openai_client:
            raise ValueError("OpenAI client not initialized. Check your API key.")
        
        # Choose appropriate prompt based on purpose
        if purpose == "chunk":
            system_prompt = f"Summarize this section concisely in {max_words} words or less. Focus on key points and important details."
        elif purpose == "final":
            system_prompt = f"Create a comprehensive summary from these section summaries. Maximum {max_words} words. Use clear, accessible language and focus on main themes."
        else:  # main
            system_prompt = f"Summarize the following document concisely. Focus on main points, key findings, and important details. Maximum {max_words} words. Use clear, accessible language."
        
        max_retries = 3
        base_delay = 1
        
        for attempt in range(max_retries):
            try:
                response = self.openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": text}
                    ],
                    max_tokens=max_words * 2,  # Rough estimate
                    temperature=0.7
                )
                
                summary = response.choices[0].message.content.strip()
                
                # Calculate cost (gpt-4o-mini pricing: $0.150/1M input, $0.600/1M output)
                input_tokens = response.usage.prompt_tokens
                output_tokens = response.usage.completion_tokens
                cost = (input_tokens * 0.150 / 1_000_000) + (output_tokens * 0.600 / 1_000_000)
                self.session_cost += cost
                
                return summary, cost
                
            except Exception as e:
                error_msg = str(e)
                
                if "rate_limit" in error_msg.lower():
                    if attempt < max_retries - 1:
                        delay = base_delay * (2 ** attempt)
                        print(f"Rate limit hit. Retrying in {delay} seconds...")
                        time.sleep(delay)
                        continue
                    else:
                        raise ValueError("Rate limit exceeded. Please try again later.")
                elif "invalid" in error_msg.lower() and "key" in error_msg.lower():
                    raise ValueError("Invalid API key. Please check your OPENAI_API_KEY.")
                else:
                    raise ValueError(f"OpenAI API error: {error_msg}")
        
        raise ValueError("Failed to generate summary after multiple attempts.")

    def summarize_doc(self, doc_id, max_words=600):
        """Generate AI summary for a document.
        
        Creates a comprehensive AI-generated summary of a document. For large
        documents (>10,000 words), uses a chunking strategy: summarizes sections
        first, then synthesizes those summaries into a final summary. Saves the
        summary to document metadata.
        
        Args:
            doc_id: Document ID.
            max_words: Maximum words for the final summary. Defaults to 600.
        
        Returns:
            tuple: (summary (str), total_cost (float)) containing the generated
                summary and total API cost.
        
        Raises:
            ValueError: If document not found or OpenAI client not initialized.
        """
        doc = self.get_doc(doc_id)
        if not doc:
            raise ValueError(f"Document with ID {doc_id} not found.")
        
        if not self.openai_client:
            raise ValueError("OpenAI client not initialized. Check your OPENAI_API_KEY.")
        
        # Show document info
        print(f"Analyzing {doc['title']} ({doc['page_count']} ", end='')
        if doc['extension'] == '.pdf':
            print("pages)...")
        elif doc['extension'] == '.docx':
            print("paragraphs)...")
        else:
            print("lines)...")
        
        # Extract full text with progress
        print("Extracting text...")
        text, word_count = self.extract_text(doc_id)
        
        total_cost = 0
        
        # Decide on summarization strategy
        if word_count < 10000:
            # Summarize all at once
            print("Generating summary...")
            summary, cost = self._summarize_text(text, max_words, purpose="main")
            total_cost += cost
        else:
            # Chunk and summarize
            num_chunks = (word_count // 3000) + 1
            print(f"Document is large ({word_count} words). Processing in {num_chunks} chunks...")
            
            chunks = self._chunk_text(text, chunk_size=3000)
            chunk_summaries = []
            
            words_per_chunk = max(50, max_words // len(chunks))
            
            for i, chunk in enumerate(chunks, 1):
                print(f"Analyzing chunk {i}/{len(chunks)}...")
                chunk_summary, cost = self._summarize_text(chunk, words_per_chunk, purpose="chunk")
                chunk_summaries.append(chunk_summary)
                total_cost += cost
            
            # Combine chunk summaries
            print("Generating final summary...")
            combined = "\n\n".join(chunk_summaries)
            summary, cost = self._summarize_text(combined, max_words, purpose="final")
            total_cost += cost
        
        # Save summary to metadata
        doc['summary'] = summary
        doc['summary_word_count'] = len(summary.split())
        self._save_documents()
        
        return summary, total_cost

    def _register_commands(self):
        """Register document-related commands.
        
        Registers all document management commands (docs-add, docs-list,
        docs-remove, docs-view, docs-extract, docs-search, docs-summarize,
        docs-refresh) with the command registry.
        """
        self.registry.register_command('docs-add', self.cmd_docs_add, 'Add a document (PDF, DOCX, TXT)', 'docs')
        self.registry.register_command('docs-list', self.cmd_docs_list, 'List all documents', 'docs')
        self.registry.register_command('docs-remove', self.cmd_docs_remove, 'Remove a document', 'docs')
        self.registry.register_command('docs-view', self.cmd_docs_view, 'View document details', 'docs')
        self.registry.register_command('docs-extract', self.cmd_docs_extract, 'Extract text from a document', 'docs')
        self.registry.register_command('docs-search', self.cmd_docs_search, 'Search across all documents', 'docs')
        self.registry.register_command('docs-summarize', self.cmd_docs_summarize, 'Generate AI summary for a document', 'docs')
        self.registry.register_command('docs-refresh', self.cmd_docs_refresh, 'Refresh document titles with intelligent extraction', 'docs')

    def cmd_docs_add(self, *args):
        """Command to add a document."""
        if not args:
            print("Error: File path is required.")
            print("Usage: add <filepath>")
            return
        
        filepath = " ".join(args)
        
        # Expand user home directory if present
        filepath = os.path.expanduser(filepath)
        
        try:
            print("Processing file...")
            doc = self.add_doc(filepath)
            
            ext_name = doc['extension'].upper().replace('.', '')
            print(f"✓ {ext_name} added: {doc['filename']} #{doc['id']}")
            print(f"  Title: {doc['title']}")
            if doc['extension'] == '.pdf':
                print(f"  Pages: {doc['page_count']}")
            elif doc['extension'] == '.docx':
                print(f"  Paragraphs: {doc['page_count']}")
            else:  # .txt
                print(f"  Lines: {doc['page_count']}")
            print(f"  Stored: {os.path.dirname(doc['filepath'])}")
            
        except FileNotFoundError as e:
            print(f"Error: {e}")
        except ValueError as e:
            print(f"Error: {e}")
        except PermissionError as e:
            print(f"Error: {e}")
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_list(self, *args):
        """Command to list all documents."""
        docs = self.list_docs()
        
        if not docs:
            print("No documents available.")
            print("Tip: Use 'add <filepath>' to add documents.")
            return
        
        table = []
        for doc in docs:
            has_summary = "Yes ✓" if doc.get('summary') else "No"
            added_date = doc['added_date'].split('T')[0]  # Just the date part
            ext_type = doc['extension'].upper().replace('.', '')
            
            # Format page count based on type
            if doc['extension'] == '.pdf':
                page_info = f"{doc['page_count']} pages"
            elif doc['extension'] == '.docx':
                page_info = f"{doc['page_count']} paras"
            else:  # .txt
                page_info = f"{doc['page_count']} lines"
            
            table.append([
                doc['id'],
                doc['title'][:40] + "..." if len(doc['title']) > 40 else doc['title'],
                ext_type,
                page_info,
                added_date,
                has_summary
            ])
        
        headers = ["ID", "Title", "Type", "Size", "Added", "Summary"]
        separator = ["─" * 2, "─" * 40, "─" * 4, "─" * 12, "─" * 10, "─" * 7]
        table.insert(0, separator)
        
        print(tabulate(table, headers=headers, tablefmt="plain"))

    def cmd_docs_remove(self, *args):
        """Command to remove a document."""
        if not args:
            print("Error: Document ID is required.")
            print("Usage: remove <doc_id>")
            return
        
        doc_id = args[0]
        
        try:
            self.remove_doc(doc_id)
                
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_view(self, *args):
        """Command to view document details."""
        if not args:
            print("Error: Document ID is required.")
            print("Usage: view <doc_id>")
            return
        
        doc_id = args[0]
        
        try:
            doc = self.get_doc(doc_id)
            if not doc:
                print(f"Error: Document with ID {doc_id} not found.")
                return
            
            # Update last accessed
            self.update_last_accessed(doc_id)
            
            # Format dates
            added_date = doc['added_date'].replace('T', ' ')
            accessed_date = doc['last_accessed'].replace('T', ' ')
            
            # Format page count
            if doc['extension'] == '.pdf':
                page_label = "Pages"
            elif doc['extension'] == '.docx':
                page_label = "Paragraphs"
            else:  # .txt
                page_label = "Lines"
            
            print("━" * 40)
            print(f"Document: {doc['title']}")
            print("━" * 40)
            print(f"ID:       {doc['id']}")
            print(f"Filename: {doc['filename']}")
            print(f"Type:     {doc['extension'].upper().replace('.', '')}")
            print(f"{page_label}: {doc['page_count']}")
            print(f"Added:    {added_date}")
            print(f"Accessed: {accessed_date}")
            
            # Show summary if exists
            if doc.get('summary'):
                print("\n" + "━" * 40)
                print("Summary:")
                print("━" * 40)
                print(doc['summary'])
                if doc.get('summary_word_count'):
                    print(f"\n({doc['summary_word_count']} words)")
            else:
                print("\nNo summary available.")
                print("Tip: Use 'docs-summarize {}' to generate a summary.".format(doc_id))
            
            print("\n" + "━" * 40)
            print("Preview:")
            print("━" * 40)
            preview = doc.get('preview', '')
            if preview:
                print(preview[:500])
            else:
                print("(No preview available)")
            
            print("━" * 40)
            
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_extract(self, *args):
        """Command to extract text from a document."""
        if not args:
            print("Error: Document ID is required.")
            print("Usage: extract <doc_id> [--page N]")
            return
        
        doc_id = args[0]
        page_num = None
        
        # Parse --page flag
        if len(args) >= 3 and args[1] == '--page':
            try:
                page_num = int(args[2])
            except ValueError:
                print("Error: Page number must be an integer.")
                return
        
        try:
            doc = self.get_doc(doc_id)
            if not doc:
                print(f"Error: Document with ID {doc_id} not found.")
                return
            
            # Show extraction message
            if page_num:
                print(f"Extracting page {page_num} from {doc['title']}...")
            else:
                print(f"Extracting full text from {doc['title']}...")
            
            # Extract text
            text, word_count = self.extract_text(doc_id, page_num)
            
            print(f"✓ Text extracted and cached ({word_count} words)\n")
            
            # Display extracted text
            print("━" * 40)
            if page_num:
                print(f"Page {page_num}")
            else:
                print(f"{doc['title']} - Full Text")
            print("━" * 40)
            print(text)
            print("━" * 40)
            
        except ValueError as e:
            print(f"Error: {e}")
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_search(self, *args):
        """Command to search across all documents."""
        if not args:
            print("Error: Search query is required.")
            print("Usage: search <query>")
            return
        
        query = " ".join(args)
        
        try:
            results = self.search_docs(query)
            
            if not results:
                print(f"No matches found for '{query}'.")
                return
            
            # Group results by document
            docs_with_matches = {}
            for doc, page_num, context in results:
                if doc['id'] not in docs_with_matches:
                    docs_with_matches[doc['id']] = {
                        'doc': doc,
                        'matches': []
                    }
                docs_with_matches[doc['id']]['matches'].append((page_num, context))
            
            print(f"\nFound in {len(docs_with_matches)} document(s):\n")
            
            for doc_id, data in docs_with_matches.items():
                doc = data['doc']
                matches = data['matches']
                
                # Show first match for each document
                page_num, context = matches[0]
                
                if page_num:
                    print(f"{doc['id']}: {doc['title']} (page {page_num})")
                else:
                    print(f"{doc['id']}: {doc['title']}")
                
                # Clean up context for display
                context_lines = context.replace('\n', ' ').strip()
                print(f"{context_lines}\n")
                
                # Show count if multiple matches
                if len(matches) > 1:
                    print(f"  (+{len(matches) - 1} more match(es) in this document)\n")
            
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_summarize(self, *args):
        """Command to generate AI summary for a document."""
        if not args:
            print("Error: Document ID is required.")
            print("Usage: summarize <doc_id> [--max-words N]")
            return
        
        doc_id = args[0]
        max_words = 600  # Default
        
        # Parse --max-words flag
        if len(args) >= 3 and args[1] == '--max-words':
            try:
                max_words = int(args[2])
                if max_words < 50:
                    print("Error: max-words must be at least 50.")
                    return
                if max_words > 2000:
                    print("Error: max-words cannot exceed 2000.")
                    return
            except ValueError:
                print("Error: max-words must be an integer.")
                return
        
        try:
            doc = self.get_doc(doc_id)
            if not doc:
                print(f"Error: Document with ID {doc_id} not found.")
                return
            
            # Check if OpenAI client is available
            if not self.openai_client:
                print("Error: OpenAI client not initialized. Check your OPENAI_API_KEY.")
                return
            
            # Generate summary
            summary, total_cost = self.summarize_doc(doc_id, max_words)
            
            # Display result
            summary_word_count = len(summary.split())
            print(f"✓ Summary created ({summary_word_count} words, estimated cost: ${total_cost:.4f})\n")
            
            print("━" * 40)
            print(f"Summary of {doc['title']}:")
            print("━" * 40)
            print(summary)
            print("━" * 40)
            
        except ValueError as e:
            print(f"Error: {e}")
        except Exception as e:
            print(f"Error: {e}")

    def cmd_docs_refresh(self, *args):
        """Command to refresh document titles with intelligent extraction."""
        if not self.documents:
            print("No documents found in the system.")
            return
        
        print(f"Refreshing titles for {len(self.documents)} document(s)...\n")
        
        updated_count = 0
        
        for doc in self.documents:
            old_title = doc['title']
            filepath = doc['filepath']
            extension = doc['extension']
            
            # Check if title needs updating (is generic)
            needs_update = False
            
            if extension == '.pdf' and old_title.lower() in ['(anonymous)', 'untitled', 'document', 'pdf']:
                needs_update = True
            elif extension == '.docx' and old_title.lower() in ['word document', 'document', 'untitled', 'doc', 'docx']:
                needs_update = True
            elif extension == '.txt' and (not old_title or old_title.isdigit()):
                needs_update = True
            
            if not needs_update:
                continue
            
            # Re-extract metadata to get new title
            try:
                if extension == '.pdf':
                    metadata = self._extract_pdf_metadata(filepath)
                elif extension == '.docx':
                    metadata = self._extract_docx_metadata(filepath)
                elif extension == '.txt':
                    metadata = self._extract_txt_metadata(filepath)
                else:
                    continue
                
                new_title = metadata['title']
                
                # If still no title, use filename without extension
                if not new_title:
                    new_title = os.path.splitext(doc['filename'])[0]
                
                # Update the title
                doc['title'] = new_title
                updated_count += 1
                
                print(f"✓ {doc['id']}: '{old_title}' → '{new_title}'")
                
            except Exception as e:
                print(f"✗ {doc['id']}: Error updating title - {e}")
        
        if updated_count > 0:
            self._save_documents()
            print(f"\n✓ Successfully refreshed {updated_count} document title(s)")
        else:
            print("No documents needed title updates.")
