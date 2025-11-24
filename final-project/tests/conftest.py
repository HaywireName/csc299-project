"""
Pytest configuration and fixtures for PKMS Task Manager tests.
"""
import pytest
import os
import json
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from io import BytesIO

# Add parent directory to path for imports
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from core.storage import DataManager, JSONStorage
from core.commands import CommandRegistry
from modules.task_module import TaskManager
from modules.docs_module import DocumentManager
from modules.chat_module import ChatManager
from modules.agent_module import AgentManager
from tests.mock_openai import create_mock_openai_client


@pytest.fixture
def temp_data_dir():
    """Create a temporary data directory for testing."""
    temp_dir = tempfile.mkdtemp(prefix="pkms_test_")
    yield Path(temp_dir)
    # Cleanup
    shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def data_manager(temp_data_dir):
    """Create a DataManager with temporary directory."""
    return DataManager(temp_data_dir)


@pytest.fixture
def command_registry():
    """Create a fresh command registry."""
    return CommandRegistry()


@pytest.fixture
def mock_openai_client():
    """Create a mock OpenAI client."""
    return create_mock_openai_client()


@pytest.fixture
def task_manager(data_manager, command_registry, mock_openai_client, monkeypatch):
    """Create a TaskManager with mocked OpenAI."""
    tm = TaskManager(data_manager, command_registry)
    # Replace OpenAI client with mock
    tm.openai_client = mock_openai_client
    # Mock environment variable
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test-key-123")
    return tm


@pytest.fixture
def document_manager(data_manager, command_registry, mock_openai_client, temp_data_dir):
    """Create a DocumentManager with mocked OpenAI."""
    # Create docs directories
    docs_dir = temp_data_dir / "docs"
    docs_dir.mkdir(exist_ok=True)
    (docs_dir / "pdfs").mkdir(exist_ok=True)
    (docs_dir / "docx").mkdir(exist_ok=True)
    (docs_dir / "txt").mkdir(exist_ok=True)
    
    # Create doc_cache directory
    cache_dir = temp_data_dir / "doc_cache"
    cache_dir.mkdir(exist_ok=True)
    
    dm = DocumentManager(data_manager, command_registry)
    # Override directories
    dm.data_dir = str(docs_dir)
    dm.pdfs_dir = str(docs_dir / "pdfs")
    dm.docx_dir = str(docs_dir / "docx")
    dm.txt_dir = str(docs_dir / "txt")
    dm.cache_dir = str(cache_dir)
    # Replace OpenAI client with mock
    dm.openai_client = mock_openai_client
    return dm


@pytest.fixture
def chat_manager(data_manager, command_registry, mock_openai_client):
    """Create a ChatManager with mocked OpenAI."""
    cm = ChatManager(data_manager, command_registry)
    # Replace OpenAI client with mock
    cm.openai_client = mock_openai_client
    return cm


@pytest.fixture
def agent_manager(data_manager, task_manager, document_manager, command_registry, mock_openai_client):
    """Create an AgentManager with mocked OpenAI."""
    am = AgentManager(data_manager, task_manager, command_registry, document_manager)
    # Replace OpenAI client with mock
    am.openai_client = mock_openai_client
    return am


@pytest.fixture
def sample_task():
    """Create a sample task dictionary."""
    return {
        "id": "1",
        "title": "Test Task",
        "description": "This is a test task description",
        "deadline": "31-12-2025",
        "priority": "medium",
        "status": "pending",
        "summary": None,
        "created": datetime.now().strftime("%d-%m-%YT%H:%M:%S")
    }


@pytest.fixture
def sample_tasks():
    """Create multiple sample tasks."""
    return [
        {
            "id": "1",
            "title": "First Task",
            "description": "Complete the project report",
            "deadline": "15-12-2025",
            "priority": "high",
            "status": "pending",
            "summary": None,
            "created": "01-12-2025T10:00:00"
        },
        {
            "id": "2",
            "title": "Second Task",
            "description": "Review the documentation",
            "deadline": "20-12-2025",
            "priority": "medium",
            "status": "pending",
            "summary": None,
            "created": "02-12-2025T11:00:00"
        },
        {
            "id": "3",
            "title": "Third Task",
            "description": "Update the tests",
            "deadline": None,
            "priority": "low",
            "status": "completed",
            "summary": None,
            "created": "03-12-2025T12:00:00"
        }
    ]


@pytest.fixture
def sample_pdf_file(temp_data_dir):
    """Create a sample PDF file for testing."""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        pdf_path = temp_data_dir / "sample.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Sample PDF Document")
        c.drawString(100, 730, "This is a test PDF file for the PKMS system.")
        c.drawString(100, 710, "It contains multiple lines of text.")
        c.showPage()
        c.drawString(100, 750, "Page 2 content")
        c.save()
        
        return str(pdf_path)
    except ImportError:
        # If reportlab is not available, create a mock PDF structure
        pytest.skip("reportlab not available for PDF generation")


@pytest.fixture
def sample_docx_file(temp_data_dir):
    """Create a sample DOCX file for testing."""
    try:
        from docx import Document
        
        docx_path = temp_data_dir / "sample.docx"
        doc = Document()
        doc.add_heading('Sample DOCX Document', 0)
        doc.add_paragraph('This is a test DOCX file for the PKMS system.')
        doc.add_paragraph('It contains multiple paragraphs of text.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Content for section 1.')
        doc.save(str(docx_path))
        
        return str(docx_path)
    except ImportError:
        pytest.skip("python-docx not available for DOCX generation")


@pytest.fixture
def sample_txt_file(temp_data_dir):
    """Create a sample TXT file for testing."""
    txt_path = temp_data_dir / "sample.txt"
    content = """Sample Text Document
This is a test TXT file for the PKMS system.
It contains multiple lines of text.

Section 1
Content for section 1.
More content here.

Section 2
Additional content.
"""
    txt_path.write_text(content)
    return str(txt_path)


@pytest.fixture
def corrupted_json_file(temp_data_dir):
    """Create a corrupted JSON file for testing error handling."""
    json_path = temp_data_dir / "corrupted.json"
    json_path.write_text("{invalid json content")
    return str(json_path)


@pytest.fixture
def sample_chat_conversation():
    """Create a sample chat conversation."""
    return {
        "id": "conv_12345678",
        "started": "2025-12-01T10:00:00",
        "messages": [
            {
                "role": "user",
                "content": "What tasks do I have?",
                "timestamp": "2025-12-01T10:00:00"
            },
            {
                "role": "assistant",
                "content": "You have 3 tasks in your default folder.",
                "timestamp": "2025-12-01T10:00:05"
            }
        ]
    }


@pytest.fixture
def long_description():
    """Create a long description for testing summarization (>100 words)."""
    return """This is a very long task description that contains more than one hundred words 
    to test the summarization feature of the task manager. The description should be comprehensive 
    enough to require AI summarization. It discusses multiple aspects of the project including 
    planning, implementation, testing, and deployment phases. Each phase has specific requirements 
    and deliverables that need to be tracked and managed. The planning phase involves gathering 
    requirements from stakeholders and creating a detailed project plan. The implementation phase 
    includes coding, code reviews, and integration testing. The testing phase covers unit tests, 
    integration tests, and end-to-end testing scenarios. Finally, the deployment phase handles 
    the release process, monitoring, and post-deployment support. All these phases need to be 
    coordinated and tracked effectively to ensure project success."""


@pytest.fixture(autouse=True)
def reset_env(monkeypatch):
    """Reset environment variables before each test."""
    monkeypatch.setenv("OPENAI_API_KEY", "sk-test-key-123")


@pytest.fixture
def capture_output(monkeypatch):
    """Capture stdout for testing print statements."""
    from io import StringIO
    output = StringIO()
    
    def mock_print(*args, **kwargs):
        print(*args, file=output, **kwargs)
    
    monkeypatch.setattr('builtins.print', mock_print)
    return output
